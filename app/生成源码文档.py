#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AIDoor 源码文档生成器
用于生成软件著作权申请的源码文档
"""

import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置信息
PROJECT_NAME = "AIDoor"
VERSION = "1.0.3"
COPYRIGHT = "著作权人：七月"
SRC_DIR = "./src"
EXTS = (".ts", ".tsx", ".js", ".jsx")

def collect_source_code():
    """收集所有源代码文件"""
    all_lines = []
    file_order = [
        # 主要入口文件
        "app.ts",
        "index.tsx", 
        "Home.tsx",
        "My.tsx",
        "layouts/index.tsx",
        "access.ts",
        "loading.tsx",
        "global.less",
        
        # 页面文件
        "pages/Home/index.tsx",
        "pages/Home/find.tsx",
        "pages/Chat/$id.tsx",
        "pages/Chat/list.tsx",
        "pages/Account/Login.tsx",
        "pages/Account/Register.tsx",
        "pages/Account/Develop.tsx",
        "pages/Detail/$id.tsx",
        "pages/My/user.tsx",
        "pages/My/develop.tsx",
        "pages/My/$type.tsx",
        "pages/My/messages.tsx",
        "pages/My/message-detail.tsx",
        "pages/My/index.tsx",
        "pages/private.tsx",
        
        # 组件文件
        "components/NavHeader/index.tsx",
        "components/FloatBtn/index.tsx",
        "components/BackNavBar/index.tsx",
        "components/VerificationCodeButton/index.tsx",
        "components/ImgUploader/index.tsx",
        "components/SystemMessages/index.tsx",
        "components/NavLink/index.tsx",
        
        # 服务文件
        "services/api/index.ts",
        "services/api/user.ts",
        "services/api/chatMessage.ts",
        "services/api/userRecord.ts",
        "services/api/publisher.ts",
        "services/api/account.ts",
        "services/api/appItem.ts",
        "services/api/banner.ts",
        "services/api/systemMessage.ts",
        "services/api/userContent.ts",
        "services/api/userFollow.ts",
        "services/api/comment.ts",
        "services/api/file.ts",
        "services/types.d.ts",
        
        # 工具文件
        "utils/imageUtils.ts",
        "utils/openUrl.ts",
        "utils/remToPx.ts",
        "utils/classNames.ts",
        "utils/format.ts",
        "utils/index.ts",
        
        # 模型文件
        "models/global.ts",
        "models/filter.ts",
        
        # 常量文件
        "constants/index.ts",
    ]
    
    # 按指定顺序收集文件
    collected_files = set()
    for file_path in file_order:
        full_path = os.path.join(SRC_DIR, file_path)
        if os.path.exists(full_path):
            try:
                with open(full_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    all_lines.append(f"// =========================================")
                    all_lines.append(f"// 文件: {file_path}")
                    all_lines.append(f"// =========================================")
                    all_lines.extend(content.split('\n'))
                    all_lines.append("")  # 空行分隔
                    collected_files.add(file_path)
            except Exception as e:
                print(f"读取文件 {file_path} 失败: {e}")
    
    # 收集其他未列出的文件
    for root, dirs, files in os.walk(SRC_DIR):
        for file in sorted(files):
            if file.endswith(EXTS):
                file_path = os.path.relpath(os.path.join(root, file), SRC_DIR)
                if file_path not in collected_files:
                    try:
                        with open(os.path.join(root, file), 'r', encoding='utf-8') as f:
                            content = f.read()
                            all_lines.append(f"// =========================================")
                            all_lines.append(f"// 文件: {file_path}")
                            all_lines.append(f"// =========================================")
                            all_lines.extend(content.split('\n'))
                            all_lines.append("")  # 空行分隔
                    except Exception as e:
                        print(f"读取文件 {file_path} 失败: {e}")
    
    return all_lines

def process_copyright_info(lines):
    """处理版权信息，替换为指定的著作权人"""
    processed_lines = []
    for line in lines:
        # 替换版权声明
        line = re.sub(r'著作权人[：:]\s*[^\n]*', f'著作权人：七月', line)
        line = re.sub(r'Copyright[^@]*@[^@]*七月[^@]*', f'Copyright © 七月', line)
        processed_lines.append(line)
    return processed_lines

def generate_docx():
    """生成DOCX文档"""
    print("=" * 50)
    print("AIDoor 源码文档生成器")
    print("=" * 50)
    
    print("正在收集源代码...")
    all_lines = collect_source_code()
    
    print("正在处理版权信息...")
    all_lines = process_copyright_info(all_lines)
    
    print(f"总代码行数: {len(all_lines)}")
    
    # 固定每页25行代码
    lines_per_page = 25
    print(f"每页行数: {lines_per_page}")
    
    # 分页处理
    pages = []
    for i in range(0, len(all_lines), lines_per_page):
        page_lines = all_lines[i:i + lines_per_page]
        pages.append(page_lines)
    
    total_actual_pages = len(pages)
    print(f"实际总页数: {total_actual_pages}")
    
    # 选择前30页和后30页
    if total_actual_pages > 60:
        selected_pages = pages[:30] + pages[-30:]
        print(f"选择前30页和后30页，共{len(selected_pages)}页")
    else:
        selected_pages = pages
        print(f"总页数不足60页，使用全部{total_actual_pages}页")
    
    # 创建DOCX文档
    print("正在生成DOCX文档...")
    doc = Document()
    
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = 'Consolas'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    style.font.size = Pt(9)
    
    # 生成页面
    for idx, page in enumerate(selected_pages, 1):
        # 添加页眉
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(f"{PROJECT_NAME} {VERSION}    {COPYRIGHT}    第{idx}页")
        header_run.font.size = Pt(10)
        header_run.font.bold = True
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加代码内容
        for line in page:
            if line.strip():  # 跳过空行
                doc.add_paragraph(line.rstrip('\n'))
        
        # 添加分页符（除了最后一页）
        if idx != len(selected_pages):
            doc.add_page_break()
    
    # 保存文档
    output_filename = f"{PROJECT_NAME}_源码文档_v{VERSION}_60页.docx"
    doc.save(output_filename)
    
    print("=" * 50)
    print("生成完成！")
    print(f"文档名称: {output_filename}")
    print(f"文档页数: {len(selected_pages)} 页")
    print(f"每页代码: {lines_per_page} 行")
    print(f"软件名称: {PROJECT_NAME}")
    print(f"版本号: {VERSION}")
    print(f"著作权人: 七月")
    print("=" * 50)

if __name__ == "__main__":
    try:
        generate_docx()
        input("按回车键退出...")
    except Exception as e:
        print(f"生成失败: {e}")
        input("按回车键退出...") 